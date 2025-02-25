from typing import List, Dict, Any
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os
import platform
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.tools import tool
from langchain_core.tools import Tool
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
import textract
import io
from PIL import Image
import numpy as np
import logging
import pdfplumber
from typing import List, Dict, Any
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from typing import Optional
import datetime
from PIL import Image, ImageDraw
from fpdf import FPDF
import spacy
from collections import defaultdict
import re

# Configuration OCR
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD", "/usr/bin/tesseract")

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            is_separator_regex=False,
        )
        self.embeddings = HuggingFaceBgeEmbeddings(
            model_name="BAAI/bge-small-en-v1.5",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Configuration logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
        # Configuration OCR basée sur le système d'exploitation
        self._configure_ocr()
        
        self.supported_plot_types = ['line', 'bar', 'scatter', 'hist', 'box', 'heatmap']
        
         # Initialize anonymization components
        self._init_anonymization()
    
    def _configure_ocr(self):
        """Configure OCR based on operating system"""
        if platform.system() == 'Windows':
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        elif platform.system() == 'Darwin':  # macOS
            pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
        else:  # Linux
            pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

    def _init_anonymization(self):
        """Initialize components needed for anonymization"""
        try:
            self.nlp_models = {
                'fr': spacy.load('fr_core_news_sm'),
                'en': spacy.load('en_core_web_sm')
            }
        except OSError:
            self.logger.info("Installing spaCy models...")
            os.system('python -m spacy download fr_core_news_sm')
            os.system('python -m spacy download en_core_web_sm')
            self.nlp_models = {
                'fr': spacy.load('fr_core_news_sm'),
                'en': spacy.load('en_core_web_sm')
            }

        self.sensitive_patterns = {
            'email': r'[\w\.-]+@[\w\.-]+\.\w+',
            'phone': r'(?:\+?\d{1,3}[-.\s]?)?\d{2,}(?:[-.\s]?\d{2,}){2,}',
            'date': r'\d{1,2}[-/]\d{1,2}[-/]\d{4}|\d{2,4}[-/]\d{1,2}[-/]\d{1,2}',
            'number': r'\b\d{6,}\b'
        }
        
        self.sensitive_field_indicators = {
            'fr': ['nom', 'prénom', 'adresse', 'téléphone', 'mail', 'passeport', 'naissance'],
            'en': ['name', 'address', 'phone', 'mail', 'passport', 'birth']
        }

    def detect_document_language(self, text: str) -> str:
        """Detect document language based on sensitive field indicators"""
        lang_scores = defaultdict(int)
        text_lower = text.lower()
        
        for lang, keywords in self.sensitive_field_indicators.items():
            for keyword in keywords:
                if keyword in text_lower:
                    lang_scores[lang] += 1
        
        return max(lang_scores.items(), key=lambda x: x[1])[0] if lang_scores else 'en'
    
    def analyze_form_structure(self, ocr_data: Dict) -> Dict:
        df = pd.DataFrame({
            'text': ocr_data['text'],
            'left': ocr_data['left'],
            'top': ocr_data['top'],
            'width': ocr_data['width'],
            'height': ocr_data['height'],
            'conf': ocr_data['conf']
        })
        
        df['line_number'] = pd.qcut(df['top'], q=20, labels=False, duplicates='drop')
        form_structure = {
            'columns': df.groupby('line_number')['left'].nunique().mean(),
            'average_line_height': df.groupby('line_number')['height'].mean().mean(),
            'label_positions': df[df['text'].str.contains(':$', regex=True)]['left'].unique()
        }
        
        return form_structure

    def is_label(self, text: str, structure: Dict) -> bool:
        if text.strip().endswith(':'):
            return True
            
        lang = self.detect_document_language(text)
        return any(indicator in text.lower() for indicator in self.sensitive_field_indicators[lang])

    def is_unknown_word(self, word: str, lang: str) -> bool:
        """Vérifie si un mot n'appartient pas au vocabulaire de SpaCy"""
        word = word.lower()
        
        if word in self.nlp_models[lang].vocab:
            return False  
        
        return True 

    def is_sensitive_value(self, text: str, prev_text: str, lang: str) -> bool:
        """Check if text contains sensitive information"""
        for pattern in self.sensitive_patterns.values():
            if re.match(pattern, text):
                return True

        doc = self.nlp_models[lang](text)
        if any(ent.label_ in ['PER', 'LOC', 'ORG', 'TIME', 'NORP', 'DATE'] for ent in doc.ents):
            return True

        if text.lower() not in self.nlp_models[lang].vocab:
            return True

        for indicator in self.sensitive_field_indicators[lang]:
            if indicator in prev_text.lower():
                return True

        return False

    def anonymize_pdf(self, input_pdf_path: str, output_pdf_path: str) -> Dict[str, Any]:
        """Anonymize PDF document"""
        try:
            images = convert_from_path(input_pdf_path)
            anonymized_images = []
            
            for image in images:
                pil_image = image.convert('RGB')
                draw = ImageDraw.Draw(pil_image)
                
                ocr_data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT, 
                                                config='--oem 3 --psm 6')
                
                form_structure = self.analyze_form_structure(ocr_data)
                page_text = " ".join(word for word in ocr_data['text'] if word.strip())
                lang = self.detect_document_language(page_text)
                
                prev_text = ""
                for i, word in enumerate(ocr_data['text']):
                    if not word.strip():
                        continue
                        
                    if self.is_label(word, form_structure):
                        prev_text = word
                        continue
                    
                    if self.is_sensitive_value(word, prev_text, lang):
                        x = ocr_data['left'][i]
                        y = ocr_data['top'][i]
                        w = ocr_data['width'][i]
                        h = ocr_data['height'][i]
                        
                        margin = 2
                        draw.rectangle([x-margin, y-margin, x+w+margin, y+h+margin], 
                                    fill="black")
                    
                    prev_text = word
                
                anonymized_images.append(pil_image)
            
            # Save anonymized PDF
            pdf = FPDF()
            temp_paths = []
            
            for i, image in enumerate(anonymized_images):
                temp_path = f"temp_image_{i}.jpg"
                temp_paths.append(temp_path)
                image.save(temp_path, "JPEG", quality=95)
                pdf.add_page()
                pdf.image(temp_path, 0, 0, 210, 297)
            
            pdf.output(output_pdf_path)
            
            # Cleanup temp files
            for temp_path in temp_paths:
                os.remove(temp_path)
            
            return {
                "status": "success",
                "message": f"PDF anonymized successfully: {output_pdf_path}",
                "anonymized_path": output_pdf_path
            }
            
        except Exception as e:
            self.logger.error(f"PDF anonymization failed: {str(e)}")
            return {
                "status": "error",
                "message": f"Failed to anonymize PDF: {str(e)}"
            }

    # Add new tool function for anonymization
    def process_document_with_anonymization(self, file_path: str) -> Dict[str, Any]:
        """Process document with anonymization"""
        try:
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}
            
            file_extension = os.path.splitext(file_path)[1].lower()
            output_path = f"anonymized_{os.path.basename(file_path)}"
            
            if file_extension == '.pdf':
                result = self.anonymize_pdf(file_path, output_path)
                return result
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file format for anonymization: {file_extension}"
                }
                
        except Exception as e:
            self.logger.error(f"Document anonymization failed: {str(e)}")
            return {"error": str(e)}

    

    def process_image_with_ocr(self, image: Image.Image) -> str:
        """Process a single image with OCR"""
        try:
            text = pytesseract.image_to_string(image, lang='fra+eng')
            return text.strip()
        except Exception as e:
            self.logger.error(f"OCR Error: {str(e)}")
            return ""

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber and OCR when needed"""
        try:
            text_content = []
            
            # Première tentative avec pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        text_content.append(text)
                    else:
                        # Si pas de texte extrait, essayer OCR
                        # Configuration pour Windows
                        pdf2image_kwargs = {}
                        if platform.system() == "Windows":
                            pdf2image_kwargs['poppler_path'] = r"C:\Program Files\poppler-23.11.0\Library\bin"
                        
                        # Convertir la page en image
                        images = convert_from_path(
                            file_path, 
                            first_page=page.page_number + 1,
                            last_page=page.page_number + 1,
                            **pdf2image_kwargs
                        )
                        
                        for image in images:
                            ocr_text = self.process_image_with_ocr(image)
                            if ocr_text:
                                text_content.append(ocr_text)
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error in PDF extraction: {str(e)}")
            
            # Fallback method: try full OCR
            try:
                self.logger.info("Attempting fallback PDF processing with full OCR...")
                pdf2image_kwargs = {}
                if platform.system() == "Windows":
                    pdf2image_kwargs['poppler_path'] = r"C:\Program Files\poppler-23.11.0\Library\bin"
                
                images = convert_from_path(file_path, **pdf2image_kwargs)
                text_content = []
                
                for image in images:
                    text = self.process_image_with_ocr(image)
                    if text:
                        text_content.append(text)
                
                return "\n\n".join(text_content)
                
            except Exception as e2:
                self.logger.error(f"Both PDF extraction methods failed. Final error: {str(e2)}")
                raise Exception(f"Could not process PDF: {str(e2)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX including text in tables"""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        content.append(row_text)
                    
            # Process any images in the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_data))
                        ocr_text = self.process_image_with_ocr(image)
                        if ocr_text:
                            content.append(ocr_text)
                    except Exception as e:
                        self.logger.error(f"Image Processing Error in DOCX: {str(e)}")
                        continue
            
            return "\n\n".join(content)
        except Exception as e:
            self.logger.error(f"DOCX Processing Error: {str(e)}")
            raise Exception(f"Could not process DOCX: {str(e)}")
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files (xlsx, xls, csv)"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert DataFrame to string representation
            buffer = io.StringIO()
            df.info(buf=buffer)
            info_str = buffer.getvalue()
            
            # Create a text representation of the data
            text_content = [
                f"Document Analysis Report for {file_path}",
                f"Number of rows: {len(df)}",
                f"Number of columns: {len(df.columns)}",
                "Columns:",
                ", ".join(df.columns.tolist()),
                "\nDataset Information:",
                info_str,
                "\nSample Data (first 5 rows):",
                df.head().to_string(),
                "\nSummary Statistics:",
                df.describe().to_string()
            ]
            
            # Store the DataFrame for later use
            self.current_df = df
            
            return "\n\n".join(text_content)
        except Exception as e:
            self.logger.error(f"Excel/CSV Processing Error: {str(e)}")
            raise Exception(f"Could not process Excel/CSV file: {str(e)}")

    def generate_plot(self, plot_type: str, x_column: str, y_column: Optional[str] = None, 
                 title: str = "", figsize: tuple = (10, 6)) -> str:
        """Generate plots from the current DataFrame
        """
        try:
            if not hasattr(self, 'current_df'):
                raise ValueError("No data loaded. Please load a data file first.")

            # Validate columns exist in DataFrame
            if x_column not in self.current_df.columns:
                raise ValueError(f"Column '{x_column}' not found in the data")
            if y_column and y_column not in self.current_df.columns:
                raise ValueError(f"Column '{y_column}' not found in the data")

            # Validate plot type
            if plot_type not in self.supported_plot_types:
                raise ValueError(f"Unsupported plot type. Supported types: {self.supported_plot_types}")
            
            # Create output directory if it doesn't exist
            output_dir = "generated_plots"
            os.makedirs(output_dir, exist_ok=True)

            # Generate unique filename based on timestamp
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{plot_type}_{timestamp}.jpg"
            filepath = os.path.join(output_dir, filename)

            # Create figure
            plt.figure(figsize=figsize)
            
            try:
                # Generate the requested plot type
                if plot_type == 'line':
                    self.current_df.plot(x=x_column, y=y_column, kind='line')
                elif plot_type == 'bar':
                    self.current_df.plot(x=x_column, y=y_column, kind='bar')
                elif plot_type == 'scatter':
                    plt.scatter(self.current_df[x_column], self.current_df[y_column])
                elif plot_type == 'hist':
                    self.current_df[x_column].hist()
                elif plot_type == 'box':
                    self.current_df.boxplot(column=x_column)
                elif plot_type == 'heatmap':
                    correlation_matrix = self.current_df.corr()
                    sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm')

                # Add labels and title
                plt.title(title)
                plt.xlabel(x_column)
                if y_column:
                    plt.ylabel(y_column)
                
                # Save plot locally
                plt.savefig(filepath, format='jpeg', bbox_inches='tight', dpi=300)

                # Save plot to base64 string
                buffer = io.BytesIO()
                plt.savefig(buffer, format='png', bbox_inches='tight')
                buffer.seek(0)
                image_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                self.logger.info(f"Plot saved successfully to {filepath}")
            
                return {
                    "base64": image_base64,
                    "file_path": filepath
                }

            except Exception as e:
                self.logger.error(f"Error during plot generation: {str(e)}")
                raise Exception(f"Failed to generate plot: {str(e)}")

            finally:
                plt.close()  # Always close the plot to free memory

        except ValueError as ve:
            self.logger.error(f"Validation error in generate_plot: {str(ve)}")
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error in generate_plot: {str(e)}")
            raise Exception(f"Plot generation failed: {str(e)}")


    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from any supported file format"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension in ['.xlsx', '.xls', '.csv']:
                return self.extract_text_from_excel(file_path)
            elif file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                # Use textract as fallback for other formats
                return textract.process(file_path).decode('utf-8')
                
        except Exception as e:
            self.logger.error(f"File Processing Error: {str(e)}")
            raise Exception(f"Could not process file: {str(e)}")

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document and return chunks with embeddings"""
        try:
            # Verify file exists
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}
            
            # Extract text
            raw_text = self.extract_text_from_file(file_path)
            if not raw_text:
                return {"error": "No text could be extracted from the document"}
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(raw_text)
            
            # Generate embeddings
            embeddings = self.embeddings.embed_documents(chunks)
            
            return {
                "text": raw_text,
                "chunks": chunks,
                "embeddings": embeddings,
                "total_chunks": len(chunks),
                "file_path": file_path
            }
        except Exception as e:
            self.logger.error(f"Document processing failed: {str(e)}")
            return {"error": str(e)}
        
    def generate_document(self, content: str, output_format: str, output_path: str) -> Dict[str, Any]:
        """Generate document in specified format (PDF, DOCX, TXT)"""
        try:
            if output_format.lower() == 'pdf':
                return self._generate_pdf(content, output_path)
            elif output_format.lower() == 'docx':
                return self._generate_docx(content, output_path)
            elif output_format.lower() == 'txt':
                return self._generate_txt(content, output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
        
        except Exception as e:
            self.logger.error(f"Document generation failed: {str(e)}")
            return {"error": str(e)}

    def _generate_pdf(self, content: str, output_path: str) -> Dict[str, Any]:
        """Generate PDF document"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Split content into lines to handle text wrapping
            lines = content.split('\n')
            for line in lines:
                # Handle long lines by wrapping text
                words = line.split()
                current_line = ''
                for word in words:
                    test_line = current_line + ' ' + word if current_line else word
                    if pdf.get_string_width(test_line) < pdf.w - 20:  # 20mm margin
                        current_line = test_line
                    else:
                        pdf.multi_cell(0, 10, current_line)
                        current_line = word
                pdf.multi_cell(0, 10, current_line)
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            pdf.output(output_path)
            
            return {
                "status": "success",
                "message": f"PDF generated successfully: {output_path}",
                "file_path": output_path
            }
        
        except Exception as e:
            raise Exception(f"Failed to generate PDF: {str(e)}")

    def _generate_docx(self, content: str, output_path: str) -> Dict[str, Any]:
        """Generate DOCX document"""
        try:
            doc = Document()
            
            # Add content to document
            paragraphs = content.split('\n\n')  # Split on double newlines for paragraphs
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            
            return {
                "status": "success",
                "message": f"DOCX generated successfully: {output_path}",
                "file_path": output_path
            }
        
        except Exception as e:
            raise Exception(f"Failed to generate DOCX: {str(e)}")

    def _generate_txt(self, content: str, output_path: str) -> Dict[str, Any]:
        """Generate TXT document"""
        try:
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "status": "success",
                "message": f"TXT file generated successfully: {output_path}",
                "file_path": output_path
            }
        
        except Exception as e:
            raise Exception(f"Failed to generate TXT file: {str(e)}")

# Add new tool for document generation
@tool
def generate_document_file(content: str, output_format: str, output_path: str) -> Dict[str, Any]:
    """Generate a document in the specified format (PDF, DOCX, TXT)"""
    processor = DocumentProcessor()
    return processor.generate_document(content, output_format, output_path)

@tool
def process_document_with_ocr(file_path: str) -> Dict[str, Any]:
    """Process a document with OCR capabilities and return extracted text and metadata"""
    processor = DocumentProcessor()
    return processor.process_document(file_path)

@tool
def generate_visualization(file_path: str, plot_type: str, x_column: str, 
                         y_column: Optional[str] = None, title: str = "") -> str:
    """Generate a visualization from the data in the specified file"""
    processor = DocumentProcessor()
    processor.process_document(file_path)
    return processor.generate_plot(plot_type, x_column, y_column, title)

# Add new tool for anonymization
@tool
def anonymize_document(file_path: str) -> Dict[str, Any]:
    """Anonymize sensitive information in documents"""
    processor = DocumentProcessor()
    return processor.process_document_with_anonymization(file_path)

# List of available tools
tools = [
    Tool(
        name="process_document_with_ocr",
        func=process_document_with_ocr,
        description="Process documents (PDF, DOCX, TXT) with OCR capabilities"
    ),
    Tool(
        name="generate_visualization",
        func=generate_visualization,
        description="Generate visualizations from data files"
    ),
    Tool(
        name="anonymize_document",
        func=anonymize_document,
        description="Anonymize sensitive information in documents"
    ),
    Tool(
        name="generate_document_file",
        func=generate_document_file,
        description="Generate documents in PDF, DOCX, or TXT format"
    )
]